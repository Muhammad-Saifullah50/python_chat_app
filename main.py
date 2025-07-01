import streamlit as st
import google.generativeai as genai
import os
from dotenv import load_dotenv
import time

# Load environment variables
load_dotenv(".env.local")
API_KEY = os.getenv("GEMINI_API_KEY")

genai.configure(api_key=API_KEY)

# Initialize chat history in session state
if "messages" not in st.session_state:
    st.session_state["messages"] = []


def get_gemini_response(messages):
    # Prepare messages for Gemini API
    chat_history = []
    for msg in messages:
        if msg["role"] == "user":
            chat_history.append({"role": "user", "parts": [msg["content"]]})
        else:
            chat_history.append({"role": "model", "parts": [msg["content"]]})
    model = genai.GenerativeModel(model_name="gemini-2.0-flash", generation_config=genai.GenerationConfig())

    chat = model.start_chat(history=chat_history)
    response = chat.send_message(messages[-1]["content"])
    return response.text


def get_gemini_response_stream(messages):
    # Prepare messages for Gemini API
    chat_history = []
    for msg in messages:
        if msg["role"] == "user":
            chat_history.append({"role": "user", "parts": [msg["content"]]})
        else:
            chat_history.append({"role": "model", "parts": [msg["content"]]})
    model = genai.GenerativeModel(model_name="gemini-2.0-flash", generation_config=genai.GenerationConfig())
    chat = model.start_chat(history=chat_history)
    # Use stream=True to get a generator of chunks
    return chat.send_message(messages[-1]["content"], stream=True)


def extract_text_from_file(uploaded_file):
    import io
    import PyPDF2
    from docx import Document
    if uploaded_file is None:
        return None
    if uploaded_file.type == "application/pdf":
        reader = PyPDF2.PdfReader(uploaded_file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text
    elif uploaded_file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
        doc = Document(uploaded_file)
        return "\n".join([p.text for p in doc.paragraphs])
    else:
        # Assume text file
        return uploaded_file.read().decode("utf-8")


def extract_text_from_files(uploaded_files):
    import io
    import PyPDF2
    from docx import Document
    all_texts = []
    for uploaded_file in uploaded_files:
        if uploaded_file.type == "application/pdf":
            reader = PyPDF2.PdfReader(uploaded_file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            all_texts.append(f"--- File: {uploaded_file.name} ---\n{text}")
        elif uploaded_file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
            doc = Document(uploaded_file)
            text = "\n".join([p.text for p in doc.paragraphs])
            all_texts.append(f"--- File: {uploaded_file.name} ---\n{text}")
        else:
            # Assume text file
            text = uploaded_file.read().decode("utf-8")
            all_texts.append(f"--- File: {uploaded_file.name} ---\n{text}")
    return "\n\n".join(all_texts)


def main():
    st.title("Python Chat App")

    # Display chat messages using st.chat_message
    for msg in st.session_state["messages"]:
        if msg["role"] == "user":
            with st.chat_message("user"):
                # Only show the user's actual query, not the file content
                if "User Query:" in msg["content"]:
                    user_display = msg["content"].split("User Query:", 1)[-1].strip()
                else:
                    user_display = msg["content"]
                st.markdown(user_display)
        else:
            with st.chat_message("assistant"):
                st.markdown(msg["content"])

    # Place multi-file uploader just above the chat input
    uploaded_files = st.file_uploader("Upload files", type=["pdf", "txt", "docx"], label_visibility="visible", key="file_uploader", accept_multiple_files=True)
    if uploaded_files:
        file_text = extract_text_from_files(uploaded_files)
        st.session_state["file_text"] = file_text
        st.session_state["last_uploaded_filenames"] = [f.name for f in uploaded_files]
    elif "file_text" not in st.session_state:
        st.session_state["file_text"] = None

    # Always show foldable file content if files are uploaded
    if st.session_state.get("file_text"):
        with st.expander("Show/hide uploaded file contents", expanded=False):
            st.code(st.session_state["file_text"][:4000], language="markdown")

    # Place chat input at the bottom
    user_input = st.chat_input("Type your message or ask about your documents...")
    if user_input:
        # If files uploaded, prepend their content to the user message for context (but only send to Gemini, not display)
        if st.session_state["file_text"]:
            user_message = f"Document Content:\n{st.session_state['file_text'][:4000]}\n\nUser Query: {user_input}"
        else:
            user_message = user_input
        st.session_state["messages"].append({"role": "user", "content": user_message})
        with st.chat_message("user"):
            st.markdown(user_input)
        with st.chat_message("assistant"):
            with st.spinner("Gemini is typing..."):
                try:
                    response_stream = get_gemini_response_stream(st.session_state["messages"])
                    full_response = ""
                    response_placeholder = st.empty()
                    animation_chars = [" ░", " ▒", " ▓", " █", " ▓", " ▒", " ░"]
                    anim_idx = 0
                    for chunk in response_stream:
                        if hasattr(chunk, 'text'):
                            for char in chunk.text:
                                full_response += char
                                # Add a blinking cursor and animated bar
                                anim = animation_chars[anim_idx % len(animation_chars)]
                                response_placeholder.markdown(full_response + f'<span style=\"color:#fcba03;font-weight:bold;\">|{anim}</span>', unsafe_allow_html=True)
                                anim_idx += 1
                                time.sleep(0.025)  # Smooth animation
                    # Finalize response without animation
                    response_placeholder.markdown(full_response)
                    st.session_state["messages"].append({"role": "model", "content": full_response})
                except Exception as e:
                    error_msg = f"Error: {e}"
                    st.session_state["messages"].append({"role": "model", "content": error_msg})
                    st.markdown(error_msg)
        st.rerun()


if __name__ == "__main__":
    main()
